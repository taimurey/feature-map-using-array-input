import numpy as np
from docx import Document

# Create a new Word document
doc = Document()

# Define the input image and filter
image = np.array([[-1, -1, -1, -1, -1, -1, -1, -1, -1],
                  [-1,  1, -1, -1, -1, -1, -1,  1, -1],
                  [-1, -1,  1, -1, -1, -1,  1, -1, -1],
                  [-1, -1, -1,  1, -1,  1, -1, -1, -1],
                  [-1, -1, -1, -1,  1, -1, -1, -1, -1],
                  [-1, -1, -1,  1, -1,  1, -1, -1, -1],
                  [-1, -1,  1, -1, -1, -1,  1, -1, -1],
                  [-1,  1, -1, -1, -1, -1, -1,  1, -1],
                  [-1, -1, -1, -1, -1, -1, -1, -1, -1]])

filter = np.array([[1, -1, -1],
                   [-1,  1, -1],
                   [-1, -1,  1]])

# Define stride and padding
stride = 1
padding = 0

# Append the input image to the document
doc.add_heading("Input Image", level=1)
table = doc.add_table(rows=len(image), cols=len(image[0]))
for i, row in enumerate(image):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        cell.text = str(value)

# Append the filter to the document
doc.add_heading("Filter", level=1)
table = doc.add_table(rows=len(filter), cols=len(filter[0]))
for i, row in enumerate(filter):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        cell.text = str(value)

# Apply the filter to generate the feature map
input_height, input_width = image.shape
filter_height, filter_width = filter.shape

output_height = (input_height - filter_height + 2 * padding) // stride + 1
output_width = (input_width - filter_width + 2 * padding) // stride + 1

doc.add_heading("Feature Map Calculation", level=1)
doc.add_paragraph(
    f"Output Height: (({input_height} - {filter_height} + 2 * {padding}) / {stride}) + 1 = {output_height}")
doc.add_paragraph(
    f"Output Width: (({input_width} - {filter_width} + 2 * {padding}) / {stride}) + 1 = {output_width}")

feature_map = np.zeros((output_height, output_width))

for i in range(output_height):
    for j in range(output_width):
        row_start = i * stride
        row_end = row_start + filter_height
        col_start = j * stride
        col_end = col_start + filter_width
        receptive_field = image[row_start:row_end, col_start:col_end]
        feature_map[i, j] = np.sum(receptive_field * filter)
        doc.add_paragraph(
            f"Feature Map[{i}, {j}] = Sum(Receptive Field * Filter) = {receptive_field} * {filter} = {feature_map[i, j]}")

# Multiply the feature map with 1/9 and round off the values
feature_map /= 9
feature_map_rounded = np.round(feature_map, 2)

doc.add_heading("Feature Map (Rounded)", level=1)
table = doc.add_table(rows=len(feature_map_rounded),
                      cols=len(feature_map_rounded[0]))
for i, row in enumerate(feature_map_rounded):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        cell.text = str(value)

# Apply max pooling operation with a 2x2 filter and stride of 2
pool_size = 2
pooled_height = output_height // pool_size
pooled_width = output_width // pool_size

doc.add_heading("Max Pooling Calculation", level=1)
doc.add_paragraph(
    f"Pooled Height: {output_height} // {pool_size} = {pooled_height}")
doc.add_paragraph(
    f"Pooled Width: {output_width} // {pool_size} = {pooled_width}")

pooled_feature_map = np.zeros((pooled_height, pooled_width))

for i in range(pooled_height):
    for j in range(pooled_width):
        row_start = i * pool_size
        row_end = row_start + pool_size
        col_start = j * pool_size
        col_end = col_start + pool_size
        pool_region = feature_map_rounded[row_start:row_end, col_start:col_end]
        pooled_feature_map[i, j] = np.max(pool_region)
        doc.add_paragraph(
            f"Pooled Feature Map[{i}, {j}] = Max(Pool Region) = {pool_region} -> {pooled_feature_map[i, j]}")

# Append the pooled feature map to the document
doc.add_heading("Pooled Feature Map", level=1)
table = doc.add_table(rows=len(pooled_feature_map),
                      cols=len(pooled_feature_map[0]))
for i, row in enumerate(pooled_feature_map):
    row_cells = table.add_row().cells
    for j, value in enumerate(row):
        row_cells[j].text = str(value)

# Save the document to a file
doc.save("output.docx")
